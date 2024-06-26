from docx import Document
from flask import render_template, request, send_file
from src import app
from src.forms import FormNome
from src.models import Nome
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import os
import time

@app.route('/', methods=['GET', 'POST'])
def home():
    form = FormNome()
    if request.method == 'POST':
        nome = request.form.get('nome')
        nome_padronizado = nome.upper()

        user = Nome.query.filter_by(username=nome_padronizado).first()
        if user:
            # Aqui ocorre a geração do documento DOCX e sua conversão para PDF
            template_path = 'CERTIFICADO_WORK.docx'
            doc = Document(template_path)

            for paragraph in doc.paragraphs:
                if 'NOME' in paragraph.text:
                    paragraph.text = paragraph.text.replace('NOME', user.username)

            pdf_path = f'temp_files\certificate_{user.username}.pdf'
            doc.save('temp_files/temp.docx')

            docx_file = 'temp_files/temp.docx'
            pdf_file = pdf_path

            # Converting DOCX to PDF
            convert_docx_to_pdf(docx_file, pdf_file)

            # Verifying if PDF file exists before sending
            if os.path.exists(pdf_file):
                return send_file(pdf_file, as_attachment=True)
            else:
                return "Arquivo PDF não encontrado", 404

    return render_template('index.html', form=form)


def convert_docx_to_pdf(input_docx, output_pdf):
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from docx import Document as DocxDocument
    from io import BytesIO

    # Read the DOCX file
    doc = DocxDocument(input_docx)

    # Create a buffer for the PDF
    buffer = BytesIO()

    # Create a PDF document
    pdf = canvas.Canvas(buffer, pagesize=letter)

    # Process each paragraph in the DOCX file
    for para in doc.paragraphs:
        text = para.text.strip()
        pdf.drawString(100, 750, text)  # Adjust position as needed

    # Save the PDF document to the output file
    pdf.save()

    # Write the buffer to the output PDF file
    with open(output_pdf, 'wb') as f:
        f.write(buffer.getvalue())
